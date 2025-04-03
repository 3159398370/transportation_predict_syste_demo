from read_data import read_data,remove_duplicates,drop_features
from missing_scale import fill_missing_values,scale_data
from keras.utils import to_categorical
from sklearn.model_selection import train_test_split
from sklearn.utils import compute_class_weight
from sklearn.tree import DecisionTreeClassifier
from sklearn.tree import export_graphviz
from sklearn.model_selection import GridSearchCV
from sklearn.metrics import accuracy_score, f1_score
import joblib
from sklearn.ensemble import RandomForestClassifier
from keras.regularizers import l1,l2
from sklearn.metrics import classification_report
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Dense
import numpy as np
import datetime
import sys
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches
from docx.shared import Inches
import matplotlib.pyplot as plt

#在此接收传入的文件地址
'''
*************你需要删除14到19行这是我的文件地址,你需要在运行此py前定义(在你的代码文件中定义，参数是共享的对吧)****************
'''
input_train_filepath = sys.argv[1]
# "D:/newSoftcup/soft-cup_1/python/data/validate_1000.csv"

# model_type = sys.argv[3] # 新增model_type变量
#input_val_filepath = r'E:\code\jupyter\3xia\A10_HOBO\data\validate_1000.csv'

#input_train_filepath是传入的训练用的训练集文件地址.
#input_train_filepath是一定有的，但是input_val_filepath(训练用的验证集）可以没有
input_file_paths = [input_train_filepath]
if 'input_val_filepath' in locals():
    input_file_paths.append(input_val_filepath)


#如果input_file_paths里文件有两个则会分别赋值给两个变量，一个则只会赋值给train_model_online_file_path
if len(input_file_paths) == 2:
    train_model_online_file_path, val_model_online_file_path = input_file_paths
else:
    train_model_online_file_path = input_file_paths[0]


#以下是对数据集进行的操作
# 判断此处传入的数据集是否包含验证集
try:
    data = read_data(train_model_online_file_path, val_model_online_file_path)
except NameError:
    data = read_data(train_model_online_file_path)

if len(data) == 4:
    df_train, y_train, df_val, y_val = data
    df_train, y_train = remove_duplicates(df_train, y_train)
    df_val = drop_features(df_val)
    df_val = fill_missing_values(df_val)
    df_val = scale_data(df_val)
else:
    df_train, y_train = data

df_train = drop_features(df_train)
df_train = fill_missing_values(df_train)
df_train = scale_data(df_train)
#到此结束 下面是模型的分析

f_score_list = []
accuracy_list = []

# 添加docx对象
doc = docx.Document()

# 添加标题
paragraph_head = doc.add_heading('算法模型评估报告')

# 设置标题格式：居中
paragraph_head_format = paragraph_head.paragraph_format
paragraph_head_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

'''
        classification = """
                      precision    recall  f1-score   support

                   0       0.84      0.89      0.86        35
                   1       0.85      0.78      0.81        36
                   2       0.81      0.74      0.77        34
                   3       0.95      1.00      0.97        35
                   4       0.94      1.00      0.97        30
                   5       0.90      0.90      0.90        30

            accuracy                           0.88       200
           macro avg       0.88      0.88      0.88       200
        weighted avg       0.88      0.88      0.88       200
        """
'''

#模型构建及训练
def train_model_online(X_train, y_train, model_type, X_val = None, y_val = None):

    if X_val is  None or y_val is  None:
        X_train, X_val, y_train, y_val = train_test_split(X_train, y_train, test_size=0.2)

    # 增加model_type判断
    model = None
    # 计算类别权重
    class_weights = compute_class_weight(class_weight='balanced',classes= np.unique(y_train),y= y_train)
    class_weights = dict(enumerate(class_weights))

    # print(sys.argv[2])
    analysis_save_path = sys.argv[2]

    # 构建模型

    if model_type == 'modelFNN':
        # 前馈神经网络模型

        doc.add_heading('前馈神经网络模型', level=2)

        print("开始训练FNN")

        model = Sequential()
        model.add(Dense(94, activation='tanh', input_shape=(X_train.shape[1],), kernel_regularizer=l1(0.01)))
        model.add(Dense(282, activation='tanh',kernel_regularizer=l2(0.01)))
        model.add(Dense(252, activation='tanh',kernel_regularizer=l2(0.01)))
        model.add(Dense(42, activation='tanh',kernel_regularizer=l2(0.01)))
        model.add(Dense(6, activation='softmax'))

        # 编译模型
        model.compile(loss='categorical_crossentropy', optimizer='Adamax', metrics=['accuracy'])

        # 将标签数据转换为独热编码
        y_train = to_categorical(y_train)
        y_val = to_categorical(y_val)

        # 训练模型
        history = model.fit(X_train,y_train, epochs= 54, batch_size= 632,class_weight=class_weights, validation_data=(X_val, y_val))

        # 在验证集上评估模型性能并输出分类报告
        y_pred = model.predict(X_val)
        y_pred = y_pred.argmax(axis=1)
        y_true = y_val.argmax(axis=1)
        classification = classification_report(y_true, y_pred)

        print("**************************FNN写入文件**************************************")
        add_Classification_table(doc,classification)

        # 计算准确率和 F 值
        accuracy = accuracy_score(y_true, y_pred)
        f1 = f1_score(y_true, y_pred, average='weighted')

        accuracy_list.append(accuracy)
        f_score_list.append(f1)
        print("FNN结束")

    if model_type == 'modelDT':
        # 决策树模型
        print("开始训练决策树模型")
        doc.add_heading('决策树模型', level=2)

        # pre_model = DecisionTreeClassifier(class_weight=class_weights)
        # 需要调优的参数范围
        # param_grid = {'max_depth': [5, 10, 15, None],
        #               'min_samples_split': [2, 5, 10],
        #               'min_samples_leaf': [1, 2, 4],
        #               'max_features': ['sqrt', 'log2']}

        # 使用 GridSearchCV 进行网格搜索和交叉验证
        # grid_search = GridSearchCV(estimator=pre_model, param_grid=param_grid, cv=5)
        # grid_search.fit(X_train, y_train)

        # 输出最佳参数
        # print("Best model parameters for modelDT: ", grid_search.best_params_)
        # doc.add_heading('最佳参数：\n'+str(grid_search.best_params_), level=6)
        # 使用最佳参数构建模型
        # {'max_depth': None, 'max_features': 'sqrt', 'min_samples_leaf': 1, 'min_samples_split': 2}
        doc.add_heading('最佳参数：\nmax_depth=None, max_features=sqrt, min_samples_leaf=1, min_samples_split=2', level=6)
        model = DecisionTreeClassifier(max_depth=None, class_weight=class_weights, max_features='sqrt', min_samples_leaf=1, min_samples_split=2)

        # 在训练集上训练最佳模型
        history = model.fit(X_train, y_train)

        # 在验证集上评估模型性能并输出分类报告
        y_pred = model.predict(X_val)
        classification = classification_report(y_val, y_pred)
        print("*****************************DT写入文件***********************************")
        add_Classification_table(doc,classification)

        # 导出决策树模型为dot文件
        # class_names = [str(c) for c in model.classes_]
        # dot_data = export_graphviz(model, out_file=None, feature_names=X_train.columns, class_names=class_names, filled=True, rounded=True)

        # 使用Graph对象将DOT格式转换为图像
        # graph = graphviz.Source(dot_data)

        # 将Graph对象保存为临时PNG文件
        # temp_file = tempfile.NamedTemporaryFile(suffix='.png')
        # temp_file = tempfile.NamedTemporaryFile(suffix='.png', dir='D:/newSoftcup/soft-cup_1/python/data')
        # graph.render(filename='./data/DT_graph', format='png')

        # 在文档中添加决策树模型的图片
        # doc.add_picture('./data/DT_graph.png', width=Inches(6))

        # 计算准确率和 F 值
        accuracy = accuracy_score(y_val, y_pred)
        f1 = f1_score(y_val, y_pred, average='weighted')

        accuracy_list.append(accuracy)
        f_score_list.append(f1)
        print("DT结束")

    if model_type == 'modelRF':
        # 随机森林模型
        print("开始训练随机森林模型")
        doc.add_heading('随机森林模型', level=2)

        # 定义超参数组合
        # param_grid = {
        #     'n_estimators': [100, 200, 300],
        #     'max_depth': [5, 10, 15],
        # }

        # 创建随机森林分类器
        # pre_model = RandomForestClassifier(class_weight=class_weights)

        # 使用GridSearchCV进行交叉验证和网格搜索
        # grid_search = GridSearchCV(estimator=pre_model, param_grid=param_grid, cv=5)
        # grid_search.fit(X_train, y_train)

        # 输出最佳参数组合和得分
        # print("Best parameters for modelRF: ", grid_search.best_params_)
        # doc.add_heading('最佳参数：\n'+str(grid_search.best_params_), level=6)

        # {'max_depth': 15, 'n_estimators': 300}
        # model = grid_search.best_estimator_
        doc.add_heading('最佳参数：\nmax_depth=15, n_estimators=300', level=6)
        model = RandomForestClassifier(max_depth=15, n_estimators=300, class_weight=class_weights)
        model.fit(X_train, y_train)
        # 在验证集上评估模型性能并输出分类报告
        y_pred = model.predict(X_val)
        classification = classification_report(y_val, y_pred)
        print("***************************RF写入文件*************************************")
        add_Classification_table(doc,classification)

        # 计算准确率和 F 值
        accuracy = accuracy_score(y_val, y_pred)
        f1 = f1_score(y_val, y_pred, average='weighted')

        accuracy_list.append(accuracy)
        f_score_list.append(f1)
        print("RF结束")

    # doc.save(analysis_save_path)
    # "D:/newSoftcup/soft-cup_1/python/test.docx"

    return analysis_save_path

def analysis_model_online(X_train, y_train, X_val = None, y_val = None):

    model_name_list = ['modelFNN','modelDT','modelRF']
    # model_name_list = ['modelRF']
    for model_type in model_name_list:
        analysis_save_path = train_model_online(X_train, y_train, model_type, X_val, y_val)

    # 绘制柱状图
    plt.bar(model_name_list, accuracy_list)

    # 添加标题和轴标签
    plt.title('Accuracy of Different Models')
    plt.xlabel('Model Name')
    plt.ylabel('Accuracy')

    # 保存图像到文件
    plt.savefig('./data/image_acc.png')

    # 绘制柱状图
    plt.bar(model_name_list, f_score_list)

    plt.title('F1_score of Different Models')
    plt.xlabel('Model Name')
    plt.ylabel('F1_score')

    # 保存图像到文件
    plt.savefig('./data/image_f1.png')

    doc.add_heading('模型准确率对比', level=2)
    doc.add_picture('./data/image_acc.png', width=Inches(4))
    doc.add_heading('模型f值对比', level=2)
    doc.add_picture('./data/image_f1.png', width=Inches(4))

    doc.save(analysis_save_path)
    return analysis_save_path

def add_Classification_table(doc,classification):
    # 添加表格并设置表格样式
    table = doc.add_table(rows=10, cols=5)
    table.style = 'Colorful Grid Accent 4'

    # 设置每一列的宽度
    table.columns[0].width = 4
    table.columns[1].width = 3
    table.columns[2].width = 3
    table.columns[3].width = 3
    table.columns[4].width = 3

    # 添加表头
    heading_cells = table.rows[0].cells
    heading_cells[0].text = "label"
    heading_cells[1].text = "Precision"
    heading_cells[2].text = "Recall"
    heading_cells[3].text = "F1-score"
    heading_cells[4].text = "support"

    lines = classification.strip().split('\n')
    print(lines)

    # 将classification报告中的内容添加到表格中
    for i in range(2, 8):
        line = lines[i].strip().split()
        # print('i=',i)
        # print(line)
        cells = table.rows[i-1].cells
        cells[0].text = line[0]
        cells[1].text = line[1]
        cells[2].text = line[2]
        cells[3].text = line[3]
        cells[4].text = line[4]


    # 将accuracy、macro avg和weighted avg添加到表格中
    for i in range(9, 10):
        line = lines[i].strip().split()
        print('i=',i)
        print(line)
        cells = table.rows[i-2].cells
        cells[0].text = line[0]
        cells[1].text = "/"
        cells[2].text = "/"
        cells[3].text = line[1]
        cells[4].text = line[2]

    for i in range(10, 12):
        line = lines[i].strip().split()
        print('i=',i)
        print(line)
        cells = table.rows[i-2].cells
        cells[0].text = line[0]+" "+line[1]
        cells[1].text = line[2]
        cells[2].text = line[3]
        cells[3].text = line[4]
        cells[4].text = line[5]

    # 设置表格文本格式
    for row in table.rows:
        for cell in row.cells:
            if not cell.paragraphs:
                cell.add_paragraph()
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                if not paragraph.alignment:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                align = paragraph.alignment
                align.vertical = WD_ALIGN_PARAGRAPH.CENTER
                align.horizontal = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = docx.shared.Pt(10)
                    font.bold = False

    # 设置表格对齐方式
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

#此变量是在线训练输出的模型文件名
if 'df_val' in locals() and 'y_val' in locals():
    output_model_analysis_name = analysis_model_online(df_train, y_train, df_val, y_val)
else:
    output_model_analysis_name = analysis_model_online(df_train, y_train)
